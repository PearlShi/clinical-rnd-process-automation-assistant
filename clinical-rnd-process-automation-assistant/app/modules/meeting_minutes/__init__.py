"""
============================================================
  meeting_minutes/ - 会议纪要自动生成模块
  临床研发流程自动化助手
============================================================
  功能说明：
  支持上传会议录音文件或会议文字转录文本，智能提取
  会议核心议题、关键决策、待执行行动项、相关责任人
  与时间要求，结构化输出标准会议纪要。
============================================================
"""

import logging
from datetime import datetime
from typing import Optional
from pathlib import Path

from app.agent.base_agent import (
    BaseAgent, TaskInput, TaskResult, MockLLMEngine
)
from app.config import MEETING_CONFIG

logger = logging.getLogger(__name__)


# ============================================================
#  会议纪要处理器
# ============================================================

class MeetingMinutesProcessor:
    """
    会议纪要处理器 - 负责核心处理逻辑。
    支持文字转录文本和音频文件两种输入方式。
    """

    def __init__(self):
        self.engine = MockLLMEngine()

    def process_transcript(self, text: str, language: str = "zh-CN") -> dict:
        """
        处理会议转录文本，生成结构化会议纪要
        Args:
            text: 会议转录文本
            language: 语言（zh-CN/en）
        Returns:
            结构化会议纪要字典
        """
        # 使用模拟引擎生成会议纪要
        result = self.engine.generate_meeting_minutes(text, language)
        return result

    def process_audio(self, audio_file_path: Path) -> tuple[bool, str, str]:
        """
        处理会议录音文件（先转文本，再处理）
        Args:
            audio_file_path: 音频文件路径
        Returns:
            (是否成功, 转录文本, 错误信息)
        """
        # 检查 speech_recognition 是否启用
        if not MEETING_CONFIG.get("enable_speech_recognition", True):
            return False, "", "语音识别功能未启用"

        try:
            import speech_recognition as sr
            from pydub import AudioSegment
        except ImportError:
            logger.warning("未安装语音识别依赖，请安装: pip install SpeechRecognition pydub")
            return False, "", "语音识别库未安装，请安装依赖: SpeechRecognition 和 pydub"

        try:
            recognizer = sr.Recognizer()

            # 转换音频为WAV格式
            audio = AudioSegment.from_file(str(audio_file_path))
            wav_path = audio_file_path.with_suffix(".wav")
            audio.export(str(wav_path), format="wav")

            # 语音识别
            with sr.AudioFile(str(wav_path)) as source:
                recognizer.adjust_for_ambient_noise(source, duration=1)
                audio_data = recognizer.record(source)

            # 尝试中文识别
            text = ""
            try:
                text = recognizer.recognize_google(audio_data, language="zh-CN")
            except:
                try:
                    text = recognizer.recognize_sphinx(audio_data, language="zh-CN")
                except:
                    text = recognizer.recognize_google(audio_data, language="en-US")

            # 清理临时文件
            if wav_path.exists():
                wav_path.unlink()

            if not text.strip():
                return False, "", "未能从音频中识别出有效文本"

            return True, text, ""

        except Exception as e:
            logger.error(f"语音识别失败: {str(e)}")
            return False, "", f"语音识别失败: {str(e)}"

    def generate_markdown(self, minutes: dict) -> str:
        """
        将结构化会议纪要转换为Markdown格式
        """
        md = []
        md.append(f"# {minutes.get('title', '会议纪要')}\n")
        md.append(f"**会议时间**: {minutes.get('meeting_time', '')}")
        md.append(f"**参会人员**: {'、'.join(minutes.get('participants', []))}\n")

        # 会议摘要
        md.append("## 会议摘要")
        md.append(minutes.get('summary', '') + "\n")

        # 核心议题
        md.append("## 核心议题")
        for i, topic in enumerate(minutes.get('core_topics', []), 1):
            md.append(f"### {i}. {topic['title']}")
            md.append(topic['content'] + "\n")

        # 关键决策
        md.append("## 关键决策")
        for i, decision in enumerate(minutes.get('key_decisions', []), 1):
            md.append(f"{i}. {decision['content']}")
        md.append("")

        # 行动项
        md.append("## 行动项")
        md.append("|序号|行动项|负责人|截止日期|优先级|")
        md.append("|---|------|------|--------|------|")
        for i, item in enumerate(minutes.get('action_items', []), 1):
            md.append(
                f"|{i}|{item['content']}|{item.get('responsible_person', '')}"
                f"|{item.get('deadline', '')}|{item.get('priority', '中')}|"
            )

        return "\n".join(md)

    def generate_docx(self, minutes: dict, output_path: Optional[Path] = None) -> Optional[Path]:
        """
        将结构化会议纪要转换为Word文档
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("未安装python-docx，无法生成Word文档")
            return None

        doc = Document()

        # 标题
        title = doc.add_heading(minutes.get('title', '会议纪要'), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = 'Light Shading Accent 1'
        info_table.cell(0, 0).text = f"会议时间: {minutes.get('meeting_time', '')}"
        info_table.cell(0, 1).text = f"参会人员: {'、'.join(minutes.get('participants', []))}"
        info_table.cell(1, 0).merge(info_table.cell(1, 1))
        info_table.cell(1, 0).text = f"摘要: {minutes.get('summary', '')}"

        doc.add_paragraph()

        # 核心议题
        doc.add_heading('核心议题', level=1)
        for i, topic in enumerate(minutes.get('core_topics', []), 1):
            doc.add_heading(f"{i}. {topic['title']}", level=2)
            doc.add_paragraph(topic['content'])

        # 关键决策
        doc.add_heading('关键决策', level=1)
        for decision in minutes.get('key_decisions', []):
            doc.add_paragraph(decision['content'], style='List Bullet')

        # 行动项
        doc.add_heading('行动项', level=1)
        action_table = doc.add_table(rows=1, cols=4)
        action_table.style = 'Light Shading Accent 1'
        hdr = action_table.rows[0].cells
        hdr[0].text = '行动项'
        hdr[1].text = '负责人'
        hdr[2].text = '截止日期'
        hdr[3].text = '优先级'
        for item in minutes.get('action_items', []):
            row = action_table.add_row().cells
            row[0].text = item['content']
            row[1].text = item.get('responsible_person', '')
            row[2].text = item.get('deadline', '')
            row[3].text = item.get('priority', '中')

        # 保存
        if output_path is None:
            output_path = Path(f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(str(output_path))
        return output_path


# ============================================================
#  会议纪要智能体
# ============================================================

class MeetingMinutesAgent(BaseAgent):
    """
    会议纪要智能体 - 处理会议转录文本和录音文件，
    自动生成结构化会议纪要。
    """

    def __init__(self):
        super().__init__(
            agent_name="会议纪要智能体",
            agent_description="自动处理会议录音和转录文本，生成结构化会议纪要"
        )
        self.processor = MeetingMinutesProcessor()

    def process(self, task_input: TaskInput) -> TaskResult:
        """
        处理会议相关输入，生成会议纪要
        支持两种输入方式：
          1. 文字转录文本（直接处理）
          2. 音频文件（先语音识别再处理）
        """
        try:
            start_time = datetime.now()

            # 判断输入类型
            if task_input.input_type == "meeting_audio" and task_input.file_path:
                # 语音文件处理
                success, transcript, error = self.processor.process_audio(task_input.file_path)
                if not success:
                    return TaskResult(
                        success=False,
                        task_type="meeting_minutes",
                        error_message=error,
                    )
                raw_text = transcript
            else:
                # 文本直接处理
                raw_text = task_input.content

            if not raw_text or not raw_text.strip():
                return TaskResult(
                    success=False,
                    task_type="meeting_minutes",
                    error_message="会议文本内容为空",
                )

            # 生成结构化会议纪要
            minutes = self.processor.process_transcript(raw_text)
            markdown = self.processor.generate_markdown(minutes)

            elapsed = (datetime.now() - start_time).total_seconds()

            return TaskResult(
                success=True,
                task_type="meeting_minutes",
                output={
                    "structured": minutes,
                    "markdown": markdown,
                    "summary": minutes.get("summary", ""),
                    "action_items": minutes.get("action_items", []),
                    "participants": minutes.get("participants", []),
                },
                raw_text=raw_text,
                processing_time=elapsed,
            )

        except Exception as e:
            logger.error(f"会议纪要处理失败: {str(e)}", exc_info=True)
            return TaskResult(
                success=False,
                task_type="meeting_minutes",
                error_message=f"处理会议纪要时出错: {str(e)}",
            )
